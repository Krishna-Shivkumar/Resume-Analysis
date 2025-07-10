import streamlit as st
import re
import ollama
import json
import spacy
import unicodedata
from docx import Document
from PyPDF2 import PdfReader
from work_exp import work_experience, work_time
from jobposting import job_info
from skills_and_education import resume_skill, extract_highest_education, extract_major
from pathlib import Path
from typing import Union
from io import BufferedReader

st.set_page_config(page_title="Resume Matcher", layout="wide")

# --- Utils ---
nlp = spacy.load("en_core_web_sm")
def extract_text(file: Union[str, Path, BufferedReader]) -> str:
    """Extract text from PDF or DOCX, supports file path or file-like object."""
    if isinstance(file, (str, Path)):
        path = Path(file)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix == ".docx":
            doc = Document(str(path))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    else:
        filename = getattr(file, "name", "").lower()
        if filename.endswith(".pdf"):
            reader = PdfReader(file)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif filename.endswith(".docx"):
            doc = Document(file)
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return ""

def clean_text(text: str) -> str:
    """Normalize and clean extracted text from PDF or DOCX."""
    text = unicodedata.normalize("NFKD", text)      # Normalize Unicode (e.g., weird quotes/spaces)
    text = text.replace('\xa0', ' ')                # Replace non-breaking spaces
    text = text.replace('\n', ' ').replace('\r', '')  # Remove newlines
    text = re.sub(r'\s+', ' ', text)                # Collapse multiple spaces
    return text.strip()

def extract_email(text: str) -> str:
    """Extract the first email address found in the text."""
    match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    return match.group(0) if match else "Not found"

def normalize_text(text: str) -> str:
    return text.lower()

def get_lemmas(text: str) -> str:
    doc = nlp(text)
    lemmas = [token.lemma_ for token in doc if not token.is_stop and token.is_alpha]
    return " ".join(lemmas)

s = extract_text('resume.pdf')
s = clean_text(s)
# print(s)
s = normalize_text(s)
print(s)
# s = get_lemmas(s)
# print(s)
print(extract_highest_education(s))