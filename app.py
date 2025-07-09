import streamlit as st
import re
import spacy
from io import BytesIO
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from google.cloud import documentai_v1 as documentai
from fpdf import FPDF
import pandas as pd

from work_exp import work_experience, work_time
from jobposting import job_info
from skills_and_education import resume_skill, extract_highest_education, extract_major

st.set_page_config(page_title="Resume Matcher", layout="wide")

# --- Google Document AI Config ---
PROJECT_ID = "resume-465414"
LOCATION = "us"
PROCESSOR_ID = "4da0edcaa19b9f53"

# --- Load Spacy model once ---
nlp = spacy.load("en_core_web_sm")

def docx_to_pdf_bytes(docx_file) -> bytes:
    doc = DocxDocument(docx_file)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            pdf.multi_cell(0, 10, text.encode("latin-1", "replace").decode("latin-1"))

    pdf.output(name="converted.pdf")
    with open("converted.pdf", "rb") as f:
        return f.read()

def parse_document_bytes(file_bytes: bytes, mime_type: str) -> str:
    client = documentai.DocumentProcessorServiceClient()
    name = f"projects/{PROJECT_ID}/locations/{LOCATION}/processors/{PROCESSOR_ID}"
    raw_document = documentai.RawDocument(content=file_bytes, mime_type=mime_type)
    request = documentai.ProcessRequest(name=name, raw_document=raw_document)

    try:
        result = client.process_document(request=request)
        return result.document.text
    except Exception as e:
        st.error(f"Error processing document with Document AI: {e}")
        return ""

def normalize_text(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9\s]", " ", text.lower())

def get_lemmas(text: str) -> str:
    doc = nlp(text)
    lemmas = [token.lemma_ for token in doc if not token.is_stop and token.is_alpha]
    return " ".join(lemmas)

def extract_email(text: str) -> str:
    match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    return match.group(0) if match else "Not found"

st.markdown("## 📊 Bulk Resume Matcher")
st.markdown("Upload one job posting and multiple resumes. We'll rank each resume against the job and provide detailed breakdowns.")

with st.form("upload_form"):
    job_file = st.file_uploader("💼 Upload Job Posting (PDF or DOCX)", type=["pdf", "docx"], key="job")
    resume_files = st.file_uploader("📄 Upload Resumes (PDFs or DOCXs)", type=["pdf", "docx"], accept_multiple_files=True, key="resumes")
    submitted = st.form_submit_button("Analyze")

if submitted and job_file and resume_files:
    with st.spinner("🔍 Extracting job posting info with Document AI..."):
        if job_file.name.lower().endswith(".docx"):
            pdf_bytes = docx_to_pdf_bytes(job_file)
            job_text = parse_document_bytes(pdf_bytes, "application/pdf")
        else:
            job_bytes = job_file.read()
            job_text = parse_document_bytes(job_bytes, "application/pdf")

        if not job_text:
            st.error("Failed to parse job posting. Please try again.")
            st.stop()

        job_text = normalize_text(job_text)
        job_text = get_lemmas(job_text)
        j_info = job_info(job_text)

    results = []
    total = len(resume_files)
    progress_bar = st.progress(0)
    EDUCATION_LEVELS = {"high school": 1, "associate": 2, "bachelor": 3, "master": 4, "phd": 5}
    job_edu = 0
    for e in EDUCATION_LEVELS.keys():
        if e in j_info.get('education_level', '').lower():
            job_edu = EDUCATION_LEVELS[e]

    for i, resume_file in enumerate(resume_files):
        with st.spinner(f"Processing resume {i+1}/{total} ({resume_file.name})..."):
            if resume_file.name.lower().endswith(".docx"):
                pdf_bytes = docx_to_pdf_bytes(resume_file)
                resume_text = parse_document_bytes(pdf_bytes, "application/pdf")
            else:
                resume_bytes = resume_file.read()
                resume_text = parse_document_bytes(resume_bytes, "application/pdf")

            if not resume_text:
                st.warning(f"Skipping {resume_file.name} due to parsing error.")
                progress_bar.progress((i + 1) / total)
                continue

            resume_text = normalize_text(resume_text)
            resume_text = get_lemmas(resume_text)

            result = resume_skill(j_info["skills"], resume_text)
            if len(result) == 2:
                score, notskills = result
                matched_skills = []
            else:
                score, notskills, matched_skills = result

            education = extract_highest_education(resume_text)
            major = extract_major(resume_text)
            work_duration = work_time(resume_text)
            temp = 0
            problems = ""
            email = extract_email(resume_text)

            for e in EDUCATION_LEVELS:
                if e in education.lower():
                    temp = EDUCATION_LEVELS[e]

            if temp >= job_edu:
                score += 100 / 3
            else:
                problems += f"User does not have the required education level. User has a <{education or 'Education Level Not Found'}> level education.\n"

            try:
                required_years = float(j_info.get("required_experience_time", 0))
            except:
                required_years = 0

            if work_duration >= required_years:
                score += 100 / 3
            elif work_duration >= 0:
                score += 100 * (work_duration / (required_years or 1))
                problems += f"User has {work_duration:.1f} years experience. Job requires {required_years} years.\n"
            else:
                problems += "Error evaluating work experience duration.\n"

            if notskills:
                problems += "Missing Skills: " + ", ".join(notskills) + "\n"

            results.append({
                "Rank": 0,
                "Email": email,
                "Score": round(score, 2),
                "Education": education or "Not Found",
                "Experience (Years)": round(work_duration, 2),
                "Matched Skills": ", ".join(matched_skills) or "None",
                "Missing Skills": ", ".join(notskills) or "None",
                "Issues": problems or "None",
                "Resume Name": resume_file.name
            })

            progress_bar.progress((i + 1) / total)

    sorted_results = sorted(results, key=lambda x: x["Score"], reverse=True)
    for idx, entry in enumerate(sorted_results):
        entry["Rank"] = idx + 1

    df = pd.DataFrame(sorted_results)

    st.markdown("## 🧾 Final Ranking Table")
    st.dataframe(df[["Rank", "Resume Name", "Email", "Score"]], use_container_width=True)

    for entry in sorted_results:
        with st.expander(f"📧 {entry['Email']} | Score: {entry['Score']} | Rank #{entry['Rank']}"):
            st.write(f"**Resume:** {entry['Resume Name']}")
            st.write(f"**Education:** {entry['Education']}")
            st.write(f"**Experience:** {entry['Experience (Years)']} years")
            st.write(f"**Matched Skills:** {entry['Matched Skills']}")
            st.write(f"**Missing Skills:** {entry['Missing Skills']}")
            st.write(f"**Issues:**\n{entry['Issues']}")

    st.success("✅ All resumes have been processed.")